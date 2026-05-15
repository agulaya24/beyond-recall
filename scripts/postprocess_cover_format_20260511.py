"""Post-process the cover-page formatting in the v11.9.9 docx.

Two cosmetic changes that are easier to apply after pandoc has produced the
docx than to express in markdown:

1. Compact the "Abstract" heading. Pandoc renders `## Abstract` as a full
   Heading2 (large bold, left-aligned). The arXiv convention is a smaller,
   centered, bold label.

2. Compact and center the abstract body paragraph. Add 0.75" indents on each
   side so the abstract reads as a narrower column than the main body. Set
   center alignment and a slightly smaller font (10pt for body, matching the
   convention for indented abstracts).

Run once after `export_v11_9_11_to_docx.py` finishes. Idempotent (running it
twice produces the same output).
"""
from __future__ import annotations

import sys
from pathlib import Path

REPO = Path(__file__).resolve().parent.parent
DOCX = REPO / "docs" / "beyond_recall_v12_1_draft.docx"


def main() -> int:
    from docx import Document
    from docx.shared import Pt, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    print(f"Opening {DOCX.name}...")
    doc = Document(str(DOCX))

    # ---- TITLE (first Heading 1) ----
    # Center it and bump its font size. The title paragraph is the very first
    # Heading 1 in the document; it sits alone at the top of the cover page.
    title_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.style and p.style.name == "Heading 1":
            title_idx = i
            break

    if title_idx is not None:
        title_p = doc.paragraphs[title_idx]
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_p.runs:
            run.font.size = Pt(20)
            run.font.bold = True
        print(f"Styled title at paragraph index {title_idx}: {title_p.text[:80]}")
    else:
        print("WARNING: no Heading 1 found for title styling")

    # Find "Abstract" heading + the next paragraph (the abstract body)
    abstract_heading_idx = None
    for i, p in enumerate(doc.paragraphs):
        if (p.style and p.style.name == "Heading 2"
                and p.text.strip() == "Abstract"):
            abstract_heading_idx = i
            break

    if abstract_heading_idx is None:
        print("ERROR: Could not find 'Abstract' Heading 2 paragraph")
        return 1

    heading_p = doc.paragraphs[abstract_heading_idx]

    print(f"Found Abstract heading at paragraph index {abstract_heading_idx}")

    # Style the heading: smaller, centered, still bold
    heading_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading_p.runs:
        run.font.size = Pt(12)
        run.font.bold = True

    # Style ALL abstract body paragraphs (every paragraph between the Abstract
    # heading and the next Heading 1/2/3, or until a page-break paragraph).
    # The v12 abstract has three paragraphs; previous v11.x had a single one.
    indent_twips = 1080  # 0.75 inch (1440 twips per inch)

    paragraphs = doc.paragraphs
    body_count = 0
    for i in range(abstract_heading_idx + 1, len(paragraphs)):
        p = paragraphs[i]

        # Stop at the next heading (Introduction etc.)
        if p.style and p.style.name in ("Heading 1", "Heading 2", "Heading 3"):
            break

        # Stop at the AI-agents footnote (entirely italic) — this is the
        # cover-page footnote that should keep its own treatment, not the
        # indented-abstract styling
        if p.runs and all(run.italic for run in p.runs if run.text.strip()):
            break

        # Stop at horizontal-rule paragraph (---)
        if p.text.strip() == "---":
            break

        # Skip empty paragraphs but continue to next
        if p.text.strip() == "":
            continue

        # Apply indent + justify + 10pt to this paragraph
        pPr = p._p.get_or_add_pPr()
        for old in pPr.findall(qn("w:ind")):
            pPr.remove(old)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(indent_twips))
        ind.set(qn("w:right"), str(indent_twips))
        pPr.append(ind)

        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.size = Pt(10)
        body_count += 1
        print(f"  Styled abstract paragraph {body_count}: {p.text[:60]}...")

    print(f"Styled {body_count} abstract body paragraphs")

    doc.save(str(DOCX))
    print(f"Saved {DOCX.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
