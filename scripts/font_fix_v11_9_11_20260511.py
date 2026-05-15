"""Hard-set font + size for v11.9.11 docx.

Body / tables / headers: Times New Roman, 11pt.
Footnotes: Times New Roman, 9pt.
Run idempotently after every export.

This is necessary because pandoc's reference-docx style inheritance is
inconsistent in practice; styles defined in the reference docx don't always
propagate to runs in the output, and direct-format overrides at 10pt / 12pt
appear in the body. We force the convention here.

Cover-format (centered "Abstract" + indented body) is NOT touched; this
script runs after `postprocess_cover_format_20260511.py`, which sets the
Abstract paragraphs to 10pt explicitly.
"""
from __future__ import annotations

import sys
from pathlib import Path

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")  # type: ignore[attr-defined]
except Exception:
    pass

REPO = Path(__file__).resolve().parent.parent
DOCX = REPO / "docs" / "beyond_recall_v11_9_11_draft.docx"

BODY_FONT = "Times New Roman"
BODY_PT = 11
FOOTNOTE_PT = 9


def main() -> int:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from lxml import etree
    import zipfile
    import shutil
    import tempfile

    print(f"Opening {DOCX.name}...")
    doc = Document(str(DOCX))

    # ---- BODY ----
    # Find the abstract body paragraph to skip (set to 10pt by cover post-process).
    # The abstract sits right after the "Abstract" heading at the start.
    paragraphs = list(doc.paragraphs)
    abstract_heading_idx = None
    for i, p in enumerate(paragraphs):
        if p.text.strip() == "Abstract":
            abstract_heading_idx = i
            break

    abstract_body_idx = abstract_heading_idx + 1 if abstract_heading_idx is not None else None

    body_runs_set = 0
    for i, p in enumerate(paragraphs):
        # Skip the abstract body paragraph (handled by cover format)
        if i == abstract_body_idx:
            continue
        for run in p.runs:
            run.font.name = BODY_FONT
            run.font.size = Pt(BODY_PT)
            # Also set East Asian font to keep consistency
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = etree.SubElement(rPr, qn("w:rFonts"))
            rFonts.set(qn("w:ascii"), BODY_FONT)
            rFonts.set(qn("w:hAnsi"), BODY_FONT)
            rFonts.set(qn("w:cs"), BODY_FONT)
            rFonts.set(qn("w:eastAsia"), BODY_FONT)
            body_runs_set += 1
    print(f"  Body paragraph runs set: {body_runs_set}")

    # ---- TABLES ----
    table_runs_set = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = BODY_FONT
                        run.font.size = Pt(BODY_PT)
                        rPr = run._element.get_or_add_rPr()
                        rFonts = rPr.find(qn("w:rFonts"))
                        if rFonts is None:
                            rFonts = etree.SubElement(rPr, qn("w:rFonts"))
                        rFonts.set(qn("w:ascii"), BODY_FONT)
                        rFonts.set(qn("w:hAnsi"), BODY_FONT)
                        rFonts.set(qn("w:cs"), BODY_FONT)
                        rFonts.set(qn("w:eastAsia"), BODY_FONT)
                        table_runs_set += 1
    print(f"  Table cell runs set: {table_runs_set}")

    doc.save(str(DOCX))

    # ---- FOOTNOTES ----
    # python-docx does not expose footnotes in its high-level API.
    # We manipulate word/footnotes.xml directly via zipfile + lxml.
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W = "{" + W_NS + "}"

    with zipfile.ZipFile(str(DOCX)) as zin:
        names = zin.namelist()
        files = {n: zin.read(n) for n in names}

    footnote_runs_set = 0
    if "word/footnotes.xml" in files:
        tree = etree.fromstring(files["word/footnotes.xml"])
        for run_el in tree.iter(W + "r"):
            # Find or create rPr inside the run
            rPr = run_el.find(W + "rPr")
            if rPr is None:
                rPr = etree.SubElement(run_el, W + "rPr")
                # rPr must come first inside the run; move it
                run_el.insert(0, rPr)
            # Remove existing rFonts and sz/szCs
            for tag in ["rFonts", "sz", "szCs"]:
                for el in rPr.findall(W + tag):
                    rPr.remove(el)
            # Add desired rFonts
            rFonts = etree.SubElement(rPr, W + "rFonts")
            rFonts.set(W + "ascii", BODY_FONT)
            rFonts.set(W + "hAnsi", BODY_FONT)
            rFonts.set(W + "cs", BODY_FONT)
            rFonts.set(W + "eastAsia", BODY_FONT)
            # Re-order rFonts to come before any other children we kept
            rPr.insert(0, rFonts)
            # Add sz / szCs (9pt = 18 half-points)
            sz = etree.SubElement(rPr, W + "sz")
            sz.set(W + "val", str(FOOTNOTE_PT * 2))
            szCs = etree.SubElement(rPr, W + "szCs")
            szCs.set(W + "val", str(FOOTNOTE_PT * 2))
            footnote_runs_set += 1
        # Serialize back with original XML declaration
        new_footnotes_xml = etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)
        files["word/footnotes.xml"] = new_footnotes_xml

    print(f"  Footnote runs set (9pt): {footnote_runs_set}")

    # ---- FootnoteReference (the in-body superscript number) ----
    # Walk document.xml, find footnoteReference markers, ensure their rPr is 9pt
    # to avoid 11pt-sized superscripts that look oversize.
    if "word/document.xml" in files:
        # Easier approach: re-read with python-docx and find runs containing footnoteReference
        # Then set their size to 9pt. We do this via direct XML walk.
        doc_tree = etree.fromstring(files["word/document.xml"])
        ftn_ref_count = 0
        for run_el in doc_tree.iter(W + "r"):
            # Check if this run contains a footnoteReference
            if run_el.find(W + "footnoteReference") is not None:
                rPr = run_el.find(W + "rPr")
                if rPr is None:
                    rPr = etree.SubElement(run_el, W + "rPr")
                    run_el.insert(0, rPr)
                # Set 9pt
                for tag in ["sz", "szCs"]:
                    for el in rPr.findall(W + tag):
                        rPr.remove(el)
                sz = etree.SubElement(rPr, W + "sz")
                sz.set(W + "val", str(FOOTNOTE_PT * 2))
                szCs = etree.SubElement(rPr, W + "szCs")
                szCs.set(W + "val", str(FOOTNOTE_PT * 2))
                ftn_ref_count += 1
        files["word/document.xml"] = etree.tostring(doc_tree, xml_declaration=True, encoding="UTF-8", standalone=True)
        print(f"  Footnote-reference superscript runs set (9pt): {ftn_ref_count}")

    # Repack the docx
    tmp_path = DOCX.with_suffix(".docx.tmp")
    with zipfile.ZipFile(str(tmp_path), "w", zipfile.ZIP_DEFLATED) as zout:
        for n in names:
            zout.writestr(n, files[n])
    shutil.move(str(tmp_path), str(DOCX))
    print(f"Saved {DOCX.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
