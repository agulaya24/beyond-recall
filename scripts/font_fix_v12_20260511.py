"""Hard-set font, size, and bold for the v12 docx, style-aware.

Targets by paragraph style:

  - Heading 2 (## section like "1. Introduction"):  16pt, bold,  TNR
  - Heading 3 (### subsection like "1.1 ..."):       13pt, bold,  TNR
  - AppendixSubhead (e.g. A.1, B.2):                 13pt, bold,  TNR
  - Title (Heading 1, set by cover-format):          skip (kept at 20pt)
  - Abstract body (set by cover-format):             skip (kept at 10pt)
  - Footnotes:                                       9pt,  TNR (existing logic)
  - Footnote references (superscript):               9pt,  TNR (existing logic)
  - All other body / tables:                         11pt, TNR

CRITICAL: We iterate ALL `<w:r>` descendants of each paragraph (via lxml),
which catches runs nested inside `<w:hyperlink>` elements. Python-docx's
`p.runs` skips those, leaving cross-reference text (e.g. §4.1) rendered
in the default Hyperlink-style font, which is visually wrong against the
TNR body. Iterating lxml-level fixes this.

Runs inside an AppendixSubhead paragraph still get bold treatment so the
mini-TOC headings visually match Heading 3.

Idempotent.
"""
from __future__ import annotations

import sys
from pathlib import Path

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")  # type: ignore[attr-defined]
except Exception:
    pass

REPO = Path(__file__).resolve().parent.parent
DOCX = REPO / "docs" / "beyond_recall_v12_1_draft.docx"

BODY_FONT = "Times New Roman"
BODY_PT = 11
FOOTNOTE_PT = 9

HEADING_SIZE_PT = {
    "Heading2": 16,
    "Heading3": 13,
    "AppendixSubhead": 13,
}


def main() -> int:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from lxml import etree
    import zipfile
    import shutil

    print(f"Opening {DOCX.name}...")
    doc = Document(str(DOCX))

    paragraphs = list(doc.paragraphs)

    # ---- Find abstract body paragraphs (skip — handled by cover format) ----
    abstract_heading_idx = None
    for i, p in enumerate(paragraphs):
        if p.text.strip() == "Abstract":
            abstract_heading_idx = i
            break

    abstract_body_indices: set[int] = set()
    if abstract_heading_idx is not None:
        for j in range(abstract_heading_idx + 1, len(paragraphs)):
            p = paragraphs[j]
            if p.style and p.style.name in ("Heading 1", "Heading 2", "Heading 3"):
                break
            if p.runs and all(run.italic for run in p.runs if run.text.strip()):
                break
            if p.text.strip() == "---":
                break
            if p.text.strip() == "":
                continue
            abstract_body_indices.add(j)

    # ---- Find title (first Heading 1) — also skip ----
    title_idx = None
    for i, p in enumerate(paragraphs):
        if p.style and p.style.name == "Heading 1":
            title_idx = i
            break

    # ---- Run-styling helper ----
    def style_run(r, font: str, size_pt: int, force_bold: bool = False) -> None:
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            rPr = etree.SubElement(r, qn("w:rPr"))
            r.remove(rPr)
            r.insert(0, rPr)

        # rFonts: replace
        for old in rPr.findall(qn("w:rFonts")):
            rPr.remove(old)
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
        rFonts.set(qn("w:ascii"), font)
        rFonts.set(qn("w:hAnsi"), font)
        rFonts.set(qn("w:cs"), font)
        rFonts.set(qn("w:eastAsia"), font)

        # sz / szCs: replace
        for tag in ("sz", "szCs"):
            for el in rPr.findall(qn("w:" + tag)):
                rPr.remove(el)
        sz = etree.SubElement(rPr, qn("w:sz"))
        sz.set(qn("w:val"), str(size_pt * 2))
        szCs = etree.SubElement(rPr, qn("w:szCs"))
        szCs.set(qn("w:val"), str(size_pt * 2))

        if force_bold:
            if rPr.find(qn("w:b")) is None:
                etree.SubElement(rPr, qn("w:b"))
            if rPr.find(qn("w:bCs")) is None:
                etree.SubElement(rPr, qn("w:bCs"))

    def paragraph_style_id(p) -> str:
        pPr = p._p.find(qn("w:pPr"))
        if pPr is None:
            return ""
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is None:
            return ""
        return pStyle.get(qn("w:val")) or ""

    # ---- BODY: walk every paragraph, apply style-aware sizing ----
    body_runs_set = 0
    heading2_runs = 0
    heading3_runs = 0
    appsub_runs = 0
    hyperlink_runs_seen = 0
    for i, p in enumerate(paragraphs):
        if i in abstract_body_indices:
            continue
        if title_idx is not None and i == title_idx:
            continue

        style_id = paragraph_style_id(p)
        if style_id in HEADING_SIZE_PT:
            size = HEADING_SIZE_PT[style_id]
            bold = True
        else:
            size = BODY_PT
            bold = False

        # Iterate every <w:r> descendant — includes hyperlink runs.
        for r in p._p.iter(qn("w:r")):
            # Detect if r is inside a hyperlink (parent <w:hyperlink>)
            parent = r.getparent()
            is_hyperlink = parent is not None and parent.tag == qn("w:hyperlink")
            if is_hyperlink:
                hyperlink_runs_seen += 1
            style_run(r, BODY_FONT, size, force_bold=bold)
            body_runs_set += 1

        if style_id == "Heading2":
            heading2_runs += 1
        elif style_id == "Heading3":
            heading3_runs += 1
        elif style_id == "AppendixSubhead":
            appsub_runs += 1

    print(f"  Body paragraphs processed (total run-sets): {body_runs_set}")
    print(f"    Heading2 paragraphs: {heading2_runs}")
    print(f"    Heading3 paragraphs: {heading3_runs}")
    print(f"    AppendixSubhead paragraphs: {appsub_runs}")
    print(f"    Hyperlink runs encountered: {hyperlink_runs_seen}")

    # ---- TABLES ----
    table_runs_set = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    style_id = paragraph_style_id(p)
                    if style_id in HEADING_SIZE_PT:
                        size = HEADING_SIZE_PT[style_id]
                        bold = True
                    else:
                        size = BODY_PT
                        bold = False
                    for r in p._p.iter(qn("w:r")):
                        style_run(r, BODY_FONT, size, force_bold=bold)
                        table_runs_set += 1
    print(f"  Table cell runs set: {table_runs_set}")

    doc.save(str(DOCX))

    # ---- FOOTNOTES + footnote-reference superscripts (unchanged from v11.9.11) ----
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W = "{" + W_NS + "}"

    with zipfile.ZipFile(str(DOCX)) as zin:
        names = zin.namelist()
        files = {n: zin.read(n) for n in names}

    footnote_runs_set = 0
    if "word/footnotes.xml" in files:
        tree = etree.fromstring(files["word/footnotes.xml"])
        for run_el in tree.iter(W + "r"):
            rPr = run_el.find(W + "rPr")
            if rPr is None:
                rPr = etree.SubElement(run_el, W + "rPr")
                run_el.insert(0, rPr)
            for tag in ["rFonts", "sz", "szCs"]:
                for el in rPr.findall(W + tag):
                    rPr.remove(el)
            rFonts = etree.SubElement(rPr, W + "rFonts")
            rFonts.set(W + "ascii", BODY_FONT)
            rFonts.set(W + "hAnsi", BODY_FONT)
            rFonts.set(W + "cs", BODY_FONT)
            rFonts.set(W + "eastAsia", BODY_FONT)
            rPr.insert(0, rFonts)
            sz = etree.SubElement(rPr, W + "sz")
            sz.set(W + "val", str(FOOTNOTE_PT * 2))
            szCs = etree.SubElement(rPr, W + "szCs")
            szCs.set(W + "val", str(FOOTNOTE_PT * 2))
            footnote_runs_set += 1
        new_footnotes_xml = etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)
        files["word/footnotes.xml"] = new_footnotes_xml
    print(f"  Footnote runs set (9pt): {footnote_runs_set}")

    if "word/document.xml" in files:
        doc_tree = etree.fromstring(files["word/document.xml"])
        ftn_ref_count = 0
        for run_el in doc_tree.iter(W + "r"):
            if run_el.find(W + "footnoteReference") is not None:
                rPr = run_el.find(W + "rPr")
                if rPr is None:
                    rPr = etree.SubElement(run_el, W + "rPr")
                    run_el.insert(0, rPr)
                for tag in ["sz", "szCs"]:
                    for el in rPr.findall(W + tag):
                        rPr.remove(el)
                sz = etree.SubElement(rPr, W + "sz")
                sz.set(W + "val", str(FOOTNOTE_PT * 2))
                szCs = etree.SubElement(rPr, W + "szCs")
                szCs.set(W + "val", str(FOOTNOTE_PT * 2))
                ftn_ref_count += 1
        files["word/document.xml"] = etree.tostring(doc_tree, xml_declaration=True, encoding="UTF-8", standalone=True)
        print(f"  Footnote-reference superscript runs set (9pt): {ftn_ref_count}")

    tmp_path = DOCX.with_suffix(".docx.tmp")
    with zipfile.ZipFile(str(tmp_path), "w", zipfile.ZIP_DEFLATED) as zout:
        for n in names:
            zout.writestr(n, files[n])
    shutil.move(str(tmp_path), str(DOCX))
    print(f"Saved {DOCX.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
