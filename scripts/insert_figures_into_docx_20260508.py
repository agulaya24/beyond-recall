"""Surgically insert PNG figures into an existing .docx by finding caption-text
paragraphs and inserting image paragraphs above them. Preserves all comments,
tracked changes, and other annotations in the source docx.

Usage:
    python insert_figures_into_docx_20260508.py <input.docx> <output.docx>

Strategy:
    Pandoc converts `![alt text](path)` to a paragraph with the alt text in
    italic when the image cannot be fetched. We locate those paragraphs by
    matching the start of the alt text ("Figure 4.1:", "Figure 4.2:", etc.)
    and insert a new paragraph immediately BEFORE each such paragraph that
    contains the image. The original italic-caption paragraph is left untouched
    so any comments or revisions attached to it are preserved.

    We do not modify the existing paragraphs at all. New image paragraphs are
    added via python-docx's element manipulation directly on the underlying
    XML tree, which is the recommended way to insert content at a specific
    position without disturbing surrounding paragraphs.
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from copy import deepcopy

REPO = Path(__file__).resolve().parent.parent
FIGURES_DIR = REPO / "figures"

# Each entry: (caption-prefix-to-match, png-filename-relative-to-figures-dir)
# Caption prefix should be unique enough to identify each figure unambiguously.
FIGURE_MAP = [
    ("Figure 4.1:", "fig_4_1_gradient_scatter_v3.png"),
    ("Figure 4.2:", "fig_4_2_compression_v3.png"),
    ("Figure 4.2.1:", "fig_4_2_1_question_improvement_rates_v3.png"),
    ("Figure 4.4.1:", "fig_4_4_1_jaccard_heatmap_v1.png"),
]

# Image width in inches (typical academic paper figure width)
IMAGE_WIDTH = Inches(6.0)


def insert_paragraph_before(reference_paragraph, doc):
    """Insert a new empty paragraph immediately before the reference paragraph.
    Returns the new paragraph object.
    """
    new_para_element = deepcopy(reference_paragraph._element)
    # Clear runs/text from the copy so we have an empty paragraph
    for child in list(new_para_element):
        # Keep paragraph properties (pPr) so formatting matches; remove runs (r)
        if child.tag == qn("w:r"):
            new_para_element.remove(child)
    reference_paragraph._element.addprevious(new_para_element)
    # Find the new paragraph in the doc.paragraphs list to return as Paragraph object
    for p in doc.paragraphs:
        if p._element is new_para_element:
            return p
    return None


def find_figure_paragraphs(doc):
    """Return a list of (FIGURE_MAP_entry, paragraph) for each figure caption found."""
    matches = []
    for prefix, filename in FIGURE_MAP:
        for para in doc.paragraphs:
            text = para.text.strip()
            if text.startswith(prefix):
                matches.append(((prefix, filename), para))
                break  # Only match the first occurrence per figure
        else:
            print(f"WARNING: caption prefix not found in docx: {prefix!r}")
    return matches


def main(input_path: Path, output_path: Path) -> int:
    print(f"Opening: {input_path}")
    doc = Document(str(input_path))

    matches = find_figure_paragraphs(doc)
    print(f"Found {len(matches)} of {len(FIGURE_MAP)} figure-caption paragraphs.")

    inserted = 0
    for (prefix, filename), caption_para in matches:
        figure_path = FIGURES_DIR / filename
        if not figure_path.exists():
            print(f"WARNING: figure file missing: {figure_path}")
            continue

        # Insert a new paragraph immediately before the caption paragraph
        image_para = insert_paragraph_before(caption_para, doc)
        if image_para is None:
            print(f"ERROR: could not create insertion paragraph for {prefix!r}")
            continue

        # Add the image as a single run in the new paragraph
        run = image_para.add_run()
        run.add_picture(str(figure_path), width=IMAGE_WIDTH)
        inserted += 1
        print(f"  Inserted {filename} above caption {prefix!r}")

    print(f"\nInserted {inserted} of {len(FIGURE_MAP)} figures.")
    print(f"Saving to: {output_path}")
    doc.save(str(output_path))
    print("Done.")
    return 0


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print(f"Usage: python {sys.argv[0]} <input.docx> <output.docx>")
        sys.exit(2)
    sys.exit(main(Path(sys.argv[1]), Path(sys.argv[2])))
