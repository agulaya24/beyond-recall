"""Apply two post-walk transformations to the v11.8 with-figures docx,
preserving Aarik's existing comments, tracked changes, and highlights:

  1. Reorder Figure 4.1 block to sit immediately after the §4.1 opener
     ("The cross-subject gradient.") — picture-after-claim per Option A.
  2. Apply row shading to any per-subject results tables (matched by
     having "Band" / "Subject" header) so low-baseline rows are tinted green,
     mid-baseline rows are tinted yellow, and high-baseline (Franklin) rows
     are tinted gray, matching the §4.1 prose's color-rendered-PDF claim.

Strategy preserves comments by *moving* lxml elements (not deepcopy + delete);
table shading is added via w:shd in tcPr, leaving cell content untouched.

If the target output path is locked by Word, the script writes to a
"<output>.NEW" fallback path and reports it.

Usage:
    python post_process_v11_8_docx_20260508.py <input.docx> <output.docx>
"""
from __future__ import annotations

import sys
from pathlib import Path

# Reconfigure stdout to UTF-8 so logging that includes Δ / unicode header
# tokens does not crash on Windows cp1252 default encoding.
try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")  # type: ignore[attr-defined]
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")  # type: ignore[attr-defined]
except Exception:
    pass

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


GREEN_FILL = "D5E8D4"   # light pastel green for low-baseline (population of relevance)
YELLOW_FILL = "FFF2CC"  # light pastel yellow for mid-baseline
GRAY_FILL = "E7E7E7"    # light gray for high-baseline (Franklin)


def reorder_figure_4_1(doc) -> bool:
    """Locate the §4.1 opener, the 'Reading the gradient' paragraph, the image
    paragraph, and the caption paragraph; relocate the Reading/image/caption
    triple to sit immediately after the opener. Idempotent: if reading is
    already directly after opener, no-op.
    """
    OPENER_FRAGMENT = "less the model already knows about a subject from pretraining"
    READING_PREFIX = "Reading the gradient"
    CAPTION_PREFIX = "Figure 4.1:"

    opener_p = reading_p = caption_p = None
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if opener_p is None and OPENER_FRAGMENT in text:
            opener_p = p
        if reading_p is None and text.startswith(READING_PREFIX):
            reading_p = p
        if caption_p is None and text.startswith(CAPTION_PREFIX):
            caption_p = p

    if not (opener_p and reading_p and caption_p):
        missing = []
        if not opener_p:
            missing.append("opener")
        if not reading_p:
            missing.append("reading-the-gradient")
        if not caption_p:
            missing.append("caption")
        print(f"REORDER skipped (missing: {missing})")
        return False

    opener_el = opener_p._element
    reading_el = reading_p._element
    caption_el = caption_p._element
    image_el = caption_el.getprevious()

    if image_el is None or image_el.tag != caption_el.tag:
        print("REORDER skipped (no image paragraph immediately before caption)")
        return False

    if opener_el.getnext() is reading_el:
        print("REORDER skipped (already in target order)")
        return False

    opener_el.addnext(reading_el)
    reading_el.addnext(image_el)
    image_el.addnext(caption_el)
    print("REORDER applied: Reading + Figure 4.1 moved after §4.1 opener")
    return True


def set_cell_shading(cell, hex_color: str) -> None:
    """Add or replace w:shd inside the cell's w:tcPr to set background color."""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)

    for old_shd in tcPr.findall(qn("w:shd")):
        tcPr.remove(old_shd)

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _band_for_c5(c5_text: str) -> str | None:
    """Given a cell's text content for the C5 baseline column, classify the
    row into 'low' (≤2.0), 'mid' (2.0–3.0), or 'high' (>3.0) by the §3.4.1
    operational thresholds. Returns None if the value can't be parsed."""
    try:
        v = float(c5_text.strip().replace('—', '').replace('-', '').strip() or 'nan')
    except ValueError:
        return None
    if v != v:  # NaN
        return None
    if v <= 2.0:
        return 'low'
    if v < 3.0:
        return 'mid'
    return 'high'


_C5_HEADER_PATTERNS = (
    "c5 baseline",
    "baseline (c5)",
    "haiku c5",
    "c5 (",
)


def _is_c5_column(header_text: str) -> bool:
    """A column header that names a C5-baseline numeric. Strict match — must
    explicitly identify itself as the C5 column, not just contain 'baseline'
    (which would false-match Δ-relative-to-baseline columns)."""
    h = header_text.strip().lower()
    return any(p in h for p in _C5_HEADER_PATTERNS) or h == "c5"


def color_per_subject_tables(doc) -> int:
    """Find tables that look like §4.1-style per-subject results (Subject as
    column 0, an explicit C5-baseline numeric in column 1) and shade rows by
    the C5 value. Older layout with a Band column as col 0 is also supported.
    """
    colored_rows = 0
    matched_tables = 0
    for ti, table in enumerate(doc.tables):
        if len(table.rows) < 2:
            continue
        header_cells = [c.text.strip().lower() for c in table.rows[0].cells]
        if len(header_cells) < 2:
            continue

        c5_col = None
        # Layout A: Subject in col 0; col 1 names a C5 baseline column.
        if "subject" in header_cells[0] and _is_c5_column(header_cells[1]):
            c5_col = 1
        # Layout B (legacy): Band/Baseline group in col 0; Subject in col 1;
        # find the explicit C5 column among the rest.
        elif (("band" in header_cells[0] or "baseline" in header_cells[0])
              and "subject" in header_cells[1]):
            for idx in range(2, len(header_cells)):
                if _is_c5_column(header_cells[idx]):
                    c5_col = idx
                    break

        if c5_col is None:
            continue

        matched_tables += 1
        print(f"  matched table #{ti} (header[:4]={header_cells[:4]}, c5_col={c5_col})")

        for row in table.rows[1:]:
            if c5_col >= len(row.cells):
                continue
            band = _band_for_c5(row.cells[c5_col].text)
            if band == 'low':
                fill = GREEN_FILL
            elif band == 'mid':
                fill = YELLOW_FILL
            elif band == 'high':
                fill = GRAY_FILL
            else:
                continue
            for cell in row.cells:
                set_cell_shading(cell, fill)
            colored_rows += 1

    print(f"COLOR applied: {colored_rows} rows shaded across {matched_tables} table(s)")
    return colored_rows


def main(input_path: Path, output_path: Path) -> int:
    print(f"Opening: {input_path}")
    doc = Document(str(input_path))

    reorder_figure_4_1(doc)
    color_per_subject_tables(doc)

    print(f"Saving: {output_path}")
    try:
        doc.save(str(output_path))
        print(f"Saved: {output_path}")
        return 0
    except PermissionError:
        fallback = output_path.with_suffix(output_path.suffix + ".NEW")
        print(f"WARNING: {output_path} is locked; writing fallback to {fallback}")
        doc.save(str(fallback))
        return 2


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print(f"Usage: python {sys.argv[0]} <input.docx> <output.docx>")
        sys.exit(2)
    sys.exit(main(Path(sys.argv[1]), Path(sys.argv[2])))
